"""
create_cover_template.py — Builds cover_template.docx programmatically with docxtpl placeholders.

This script creates the exact master cover page template matching AB Ark's reference PDF layout.
It uses docxtpl Jinja2 tags for dynamic variables and InlineImage placeholders for logos.

Run directly to regenerate template:
    python -m src.services.proposal_generator.create_cover_template
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .constants import (
    HEX_PRIMARY_BLUE, RGB_PRIMARY_BLUE, RGB_TEXT_BLACK,
    MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT,
    FONT_PRIMARY, SIZE_COVER_TITLE, SIZE_COVER_SUBTITLE, SIZE_COVER_FROM,
    DEFAULT_DISCLAIMER
)


def _add_bottom_border_to_paragraph(paragraph, hex_color: str = HEX_PRIMARY_BLUE, size_eighth_pt: int = 12):
    """Adds a horizontal bottom rule directly beneath a paragraph using OpenXML."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighth_pt))  # 12 = 1.5pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_cover_template(output_path: str):
    """
    Programmatically builds cover_template.docx for docxtpl rendering.
    """
    doc = Document()
    
    # Configure margins (0.8" on all sides)
    section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Ensure section does not inherit headers/footers for cover page
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # Set default document font
    doc.styles['Normal'].font.name = FONT_PRIMARY

    # ================= 1. HEADER ROW (Above Blue Line) ================= #
    # 3-Column borderless table: Logo Left | Title Center | DocID Right
    head_table = doc.add_table(rows=1, cols=3)
    head_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_table.autofit = False

    widths = [Inches(2.0), Inches(2.9), Inches(2.0)]
    for cell, w in zip(head_table.rows[0].cells, widths):
        cell.width = w

    # Left: Header Logo placeholder
    cell_l = head_table.rows[0].cells[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_l.add_run("{{ header_logo }}")

    # Center: "Proposal (Technical)"
    cell_m = head_table.rows[0].cells[1]
    p_m = cell_m.paragraphs[0]
    p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_m = p_m.add_run("Proposal (Technical)")
    run_m.font.name = FONT_PRIMARY
    run_m.font.size = Pt(11)
    run_m.font.bold = True
    run_m.font.color.rgb = RGB_TEXT_BLACK

    # Right: Doc ID & Version Jinja2 placeholders
    cell_r = head_table.rows[0].cells[2]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r = p_r.add_run("{{ proposal_id }}\nVersion: {{ version }}")
    run_r.font.name = FONT_PRIMARY
    run_r.font.size = Pt(9.5)
    run_r.font.bold = True
    run_r.font.color.rgb = RGB_TEXT_BLACK

    # Thin blue horizontal rule near top
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_after = Pt(24)
    p_rule.paragraph_format.space_before = Pt(4)
    _add_bottom_border_to_paragraph(p_rule, hex_color=HEX_PRIMARY_BLUE, size_eighth_pt=12)

    # ================= 2. LARGE CENTERED TITLE ================= #
    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_proj.paragraph_format.space_before = Pt(12)
    p_proj.paragraph_format.space_after = Pt(18)
    run_proj = p_proj.add_run("{{ client_project_name }}")
    run_proj.font.name = FONT_PRIMARY
    run_proj.font.size = SIZE_COVER_TITLE
    run_proj.font.bold = True
    run_proj.font.color.rgb = RGB_PRIMARY_BLUE

    # ================= 3. CENTERED AB ARK LOGO (LARGE) ================= #
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(12)
    p_logo.paragraph_format.space_after = Pt(24)
    p_logo.add_run("{{ center_logo }}")

    # ================= 4. PROPOSAL SUBTITLE ================= #
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(12)
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("{{ proposal_subtitle }}")
    run_sub.font.name = FONT_PRIMARY
    run_sub.font.size = SIZE_COVER_SUBTITLE
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGB_PRIMARY_BLUE

    # ================= 5. FROM / TO BLOCK ================= #
    p_from = doc.add_paragraph()
    p_from.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_from.paragraph_format.space_before = Pt(12)
    p_from.paragraph_format.space_after = Pt(48)
    run_from = p_from.add_run("From AB Ark\nto\n{{ client_name }}")
    run_from.font.name = FONT_PRIMARY
    run_from.font.size = SIZE_COVER_FROM
    run_from.font.bold = True
    run_from.font.color.rgb = RGB_PRIMARY_BLUE

    # ================= 6. LOWER METADATA & DISCLAIMER ================= #
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(24)
    p_meta.paragraph_format.space_after = Pt(4)

    run_id_lbl = p_meta.add_run("Proposal ID: ")
    run_id_lbl.font.bold = True
    run_id_lbl.font.color.rgb = RGB_PRIMARY_BLUE
    run_id_lbl.font.size = Pt(9.5)

    run_id_val = p_meta.add_run("{{ proposal_id }}\n")
    run_id_val.font.bold = True
    run_id_val.font.color.rgb = RGB_PRIMARY_BLUE
    run_id_val.font.size = Pt(9.5)

    run_date_lbl = p_meta.add_run("Dated: ")
    run_date_lbl.font.bold = True
    run_date_lbl.font.color.rgb = RGB_PRIMARY_BLUE
    run_date_lbl.font.size = Pt(9.5)

    run_date_val = p_meta.add_run("{{ date }}")
    run_date_val.font.bold = True
    run_date_val.font.color.rgb = RGB_PRIMARY_BLUE
    run_date_val.font.size = Pt(9.5)

    # Disclaimer text
    p_disc = doc.add_paragraph()
    p_disc.paragraph_format.space_before = Pt(4)
    p_disc.paragraph_format.space_after = Pt(12)
    run_disc = p_disc.add_run(DEFAULT_DISCLAIMER)
    run_disc.font.name = FONT_PRIMARY
    run_disc.font.size = Pt(8.5)
    run_disc.font.italic = True
    run_disc.font.color.rgb = RGB_PRIMARY_BLUE

    # Save to data/templates/
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Cover template created at: '{output_path}'")


if __name__ == "__main__":
    template_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))),
        "data", "templates"
    )
    target_path = os.path.join(template_dir, "cover_template.docx")
    build_cover_template(target_path)
