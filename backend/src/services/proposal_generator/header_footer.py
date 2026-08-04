"""
header_footer.py — Header and Footer Configuration Module for python-docx.

Implements AB Ark's exact header (3-column borderless table + bottom blue rule) and
footer (top blue rule, clickable hyperlinks, 2-column layout with vertical blue divider,
and dynamic Word "Page X of Y" field codes).
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .constants import (
    FONT_PRIMARY, HEX_PRIMARY_BLUE, RGB_PRIMARY_BLUE, RGB_TEXT_BLACK, RGB_TEXT_MUTED,
    SIZE_HEADER_FOOTER, HEADER_LOGO_WIDTH, DEFAULT_FOOTER_COPYRIGHT
)
from .utils import generate_proposal_id, get_default_version


# ---------------- XML HELPERS FOR BORDERS & HYPERLINKS ---------------- #

def _add_bottom_border_to_paragraph(paragraph, hex_color: str = HEX_PRIMARY_BLUE, size_eighth_pt: int = 12):
    """Adds a thin horizontal bottom rule directly beneath a paragraph using OpenXML."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighth_pt))  # 12 = 1.5pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_top_border_to_paragraph(paragraph, hex_color: str = HEX_PRIMARY_BLUE, size_eighth_pt: int = 12):
    """Adds a thin horizontal top rule directly above a paragraph using OpenXML."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size_eighth_pt))  # 12 = 1.5pt
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), hex_color)
    pBdr.append(top)
    pPr.append(pBdr)


def _add_left_border_to_cell(cell, hex_color: str = HEX_PRIMARY_BLUE, size_eighth_pt: int = 12):
    """Adds a thin vertical left border divider to a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size_eighth_pt))  # 12 = 1.5pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    tcBorders.append(left)
    tcPr.append(tcBorders)


def add_hyperlink_to_paragraph(paragraph, url: str, text: str, hex_color: str = HEX_PRIMARY_BLUE, underline: bool = True):
    """
    Appends a clickable hyperlink to a paragraph in python-docx by registering relationship XML.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_PRIMARY)
    rFonts.set(qn('w:hAnsi'), FONT_PRIMARY)
    rPr.append(rFonts)

    # Set Font Size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(SIZE_HEADER_FOOTER.pt * 2)))
    rPr.append(sz)

    if hex_color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), hex_color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number_fields(paragraph):
    """
    Injects dynamic Word PAGE and NUMPAGES field codes into a paragraph for 'Page X of Y'.
    """
    # "Page "
    r1 = paragraph.add_run("Page ")
    r1.font.name = FONT_PRIMARY
    r1.font.size = SIZE_HEADER_FOOTER
    r1.font.italic = True
    r1.font.color.rgb = RGB_TEXT_MUTED

    # Field: PAGE
    fld1 = OxmlElement('w:fldSimple')
    fld1.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld1)

    # " of "
    r2 = paragraph.add_run(" of ")
    r2.font.name = FONT_PRIMARY
    r2.font.size = SIZE_HEADER_FOOTER
    r2.font.italic = True
    r2.font.color.rgb = RGB_TEXT_MUTED

    # Field: NUMPAGES
    fld2 = OxmlElement('w:fldSimple')
    fld2.set(qn('w:instr'), 'NUMPAGES')
    paragraph._p.append(fld2)


# ---------------- MAIN SETUP FUNCTION ---------------- #

def setup_header_and_footer(
    section,
    logo_path: str = None,
    doc_id: str = None,
    version: str = None
):
    """
    Configures headers and footers on a python-docx Section object.
    Unlinks from previous section so cover page stays clean.
    """
    if not doc_id:
        doc_id = generate_proposal_id()
    if not version:
        version = get_default_version()

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # ================= HEADER SETUP ================= #
    header = section.header
    p_head_default = header.paragraphs[0]
    p_head_default.text = ""

    # 3-column table
    head_table = header.add_table(rows=1, cols=3, width=Inches(6.9))
    head_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_table.autofit = False

    widths = [Inches(2.0), Inches(2.9), Inches(2.0)]
    for cell, w in zip(head_table.rows[0].cells, widths):
        cell.width = w

    # Left: Logo
    cell_l = head_table.rows[0].cells[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_l.paragraph_format.space_after = Pt(2)
    if logo_path and os.path.exists(logo_path):
        p_l.add_run().add_picture(logo_path, width=HEADER_LOGO_WIDTH)

    # Center: Proposal Title
    cell_m = head_table.rows[0].cells[1]
    p_m = cell_m.paragraphs[0]
    p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m.paragraph_format.space_after = Pt(2)
    r_m = p_m.add_run("Proposal (Technical)")
    r_m.font.name = FONT_PRIMARY
    r_m.font.size = Pt(10.5)
    r_m.font.bold = True
    r_m.font.color.rgb = RGB_TEXT_BLACK

    # Right: Doc ID & Version
    cell_r = head_table.rows[0].cells[2]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.paragraph_format.space_after = Pt(2)
    r_r = p_r.add_run(f"{doc_id}\nVersion: {version}")
    r_r.font.name = FONT_PRIMARY
    r_r.font.size = SIZE_HEADER_FOOTER
    r_r.font.bold = True
    r_r.font.color.rgb = RGB_TEXT_BLACK

    # Bottom blue rule beneath header table
    p_rule = header.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(2)
    p_rule.paragraph_format.space_after = Pt(12)
    _add_bottom_border_to_paragraph(p_rule, hex_color=HEX_PRIMARY_BLUE, size_eighth_pt=12)

    # ================= FOOTER SETUP ================= #
    footer = section.footer
    p_foot_default = footer.paragraphs[0]
    p_foot_default.text = ""

    # Top blue rule above footer
    _add_top_border_to_paragraph(p_foot_default, hex_color=HEX_PRIMARY_BLUE, size_eighth_pt=12)
    p_foot_default.paragraph_format.space_before = Pt(4)
    p_foot_default.paragraph_format.space_after = Pt(4)

    # Hyperlinks row (Centered)
    p_links = footer.add_paragraph()
    p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_links.paragraph_format.space_before = Pt(0)
    p_links.paragraph_format.space_after = Pt(6)

    add_hyperlink_to_paragraph(p_links, "https://www.abark.tech", "www.abark.tech")
    r_sep1 = p_links.add_run("  |  ")
    r_sep1.font.name = FONT_PRIMARY
    r_sep1.font.size = SIZE_HEADER_FOOTER
    r_sep1.font.color.rgb = RGB_TEXT_MUTED

    add_hyperlink_to_paragraph(p_links, "mailto:contact@abark.pk", "contact@abark.pk")
    r_sep2 = p_links.add_run("  |  ")
    r_sep2.font.name = FONT_PRIMARY
    r_sep2.font.size = SIZE_HEADER_FOOTER
    r_sep2.font.color.rgb = RGB_TEXT_MUTED

    add_hyperlink_to_paragraph(p_links, "tel:+923288028640", "+92 328 8028640")

    # 2-Column row with vertical blue divider for Copyright & Page Numbers
    foot_table = footer.add_table(rows=1, cols=2, width=Inches(6.9))
    foot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    foot_table.autofit = False

    w_cols = [Inches(4.5), Inches(2.4)]
    for cell, w in zip(foot_table.rows[0].cells, w_cols):
        cell.width = w

    # Left: Copyright text
    cell_f_left = foot_table.rows[0].cells[0]
    p_fl = cell_f_left.paragraphs[0]
    p_fl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_fl.paragraph_format.space_after = Pt(0)
    r_cpr = p_fl.add_run(DEFAULT_FOOTER_COPYRIGHT)
    r_cpr.font.name = FONT_PRIMARY
    r_cpr.font.size = SIZE_HEADER_FOOTER
    r_cpr.font.italic = True
    r_cpr.font.color.rgb = RGB_TEXT_MUTED

    # Right: Vertical blue divider + Page X of Y
    cell_f_right = foot_table.rows[0].cells[1]
    _add_left_border_to_cell(cell_f_right, hex_color=HEX_PRIMARY_BLUE, size_eighth_pt=12)
    p_fr = cell_f_right.paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fr.paragraph_format.space_after = Pt(0)
    add_page_number_fields(p_fr)


# Verification runner when executed directly
if __name__ == "__main__":
    doc = Document()
    setup_header_and_footer(doc.sections[0])
    print("--- header_footer.py verification completed successfully ---")
