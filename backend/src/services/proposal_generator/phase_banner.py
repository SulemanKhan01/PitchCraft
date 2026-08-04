"""
phase_banner.py — Custom Visual "Phase Banner" Component for python-docx.

Renders full-width dark navy callout boxes with orange labels, white bold title,
and vertical mini-columns for Timeline and Investment.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .constants import (
    FONT_PRIMARY, HEX_NAVY, RGB_WHITE, RGB_ORANGE,
    SIZE_BANNER_LABEL, SIZE_BANNER_TITLE, SIZE_BANNER_VALUE
)


def _set_cell_background(cell, hex_color: str):
    """Sets background shading hex color for a table cell using OpenXML."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Sets internal padding (in twips) for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_left_border_to_cell(cell, hex_color: str = "475569", size_eighth_pt: int = 8):
    """Adds a thin vertical left border divider to a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size_eighth_pt))  # 8 = 1pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    tcBorders.append(left)
    tcPr.append(tcBorders)


def add_phase_banner(
    doc: Document,
    phase_label: str,
    title: str,
    timeline: str,
    investment: str
):
    """
    Renders a dark navy visual Phase Banner table block matching the reference PDF:
    ---------------------------------------------------------------
    | PHASE 1                       | Timeline    | Investment    |
    | MVP Production Build          | Weeks 1 - 6 | $8,000        |
    ---------------------------------------------------------------
    """
    # 1-row, 3-column table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Width distribution: Total ~6.9"
    col_widths = [Inches(3.9), Inches(1.5), Inches(1.5)]
    row = table.rows[0]

    for cell, width in zip(row.cells, col_widths):
        cell.width = width
        _set_cell_background(cell, HEX_NAVY)
        _set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

    # ---------------- Cell 0: Label + Title (Left Side) ---------------- #
    cell_left = row.cells[0]
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.space_before = Pt(0)
    p_left.paragraph_format.space_after = Pt(2)
    
    # Orange Label
    r_label = p_left.add_run(f"{phase_label.upper()}\n")
    r_label.font.name = FONT_PRIMARY
    r_label.font.size = SIZE_BANNER_LABEL
    r_label.font.bold = True
    r_label.font.color.rgb = RGB_ORANGE

    # White Title
    r_title = p_left.add_run(title)
    r_title.font.name = FONT_PRIMARY
    r_title.font.size = SIZE_BANNER_TITLE
    r_title.font.bold = True
    r_title.font.color.rgb = RGB_WHITE

    # ---------------- Cell 1: Timeline (Middle Mini-Column) ---------------- #
    cell_mid = row.cells[1]
    _add_left_border_to_cell(cell_mid, hex_color="475569", size_eighth_pt=8)
    p_mid = cell_mid.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mid.paragraph_format.space_before = Pt(0)
    p_mid.paragraph_format.space_after = Pt(2)

    r_tm_lbl = p_mid.add_run("Timeline\n")
    r_tm_lbl.font.name = FONT_PRIMARY
    r_tm_lbl.font.size = Pt(8.5)
    r_tm_lbl.font.color.rgb = RGBColor(148, 163, 184)  # Muted silver/gray

    r_tm_val = p_mid.add_run(timeline)
    r_tm_val.font.name = FONT_PRIMARY
    r_tm_val.font.size = SIZE_BANNER_VALUE
    r_tm_val.font.bold = True
    r_tm_val.font.color.rgb = RGB_WHITE

    # ---------------- Cell 2: Investment (Right Mini-Column) ---------------- #
    cell_right = row.cells[2]
    _add_left_border_to_cell(cell_right, hex_color="475569", size_eighth_pt=8)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.space_before = Pt(0)
    p_right.paragraph_format.space_after = Pt(2)

    r_inv_lbl = p_right.add_run("Investment\n")
    r_inv_lbl.font.name = FONT_PRIMARY
    r_inv_lbl.font.size = Pt(8.5)
    r_inv_lbl.font.color.rgb = RGBColor(148, 163, 184)

    r_inv_val = p_right.add_run(investment)
    r_inv_val.font.name = FONT_PRIMARY
    r_inv_val.font.size = Pt(13)
    r_inv_val.font.bold = True
    r_inv_val.font.color.rgb = RGB_ORANGE

    # Add spacing after banner table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

    return table


# Verification test when run directly
if __name__ == "__main__":
    doc = Document()
    add_phase_banner(
        doc,
        phase_label="PHASE 1",
        title="MVP Production Build",
        timeline="Weeks 1 – 6",
        investment="$8,000"
    )
    print("--- phase_banner.py verification completed successfully ---")
