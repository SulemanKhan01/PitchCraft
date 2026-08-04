"""
markdown_parser.py — Dynamic Markdown-to-DOCX Parser for AB {Ark}.

Parses raw Markdown strings into python-docx elements matching AB Ark's design system:
- Heading 1 (##) -> Blue H1 with bottom underline rule
- Heading 2 (###) -> Bold black H2
- Markdown Tables (| Col1 | Col2 |) -> Styled Word tables (Navy headers + zebra rows)
- Phase Banners -> Dark navy visual callout boxes
- Bullet points (- / *) -> List Bullet style
"""

import re
from docx import Document
from .styles import add_custom_heading_1, format_data_table
from .phase_banner import add_phase_banner


def _is_table_row(line: str) -> bool:
    """Checks if a string line is a markdown table row."""
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|')


def _is_table_separator(line: str) -> bool:
    """Checks if line is a markdown table header separator (e.g. |---|---|)."""
    stripped = line.strip()
    return bool(re.match(r'^\|(?:\s*:?-+:?\s*\|)+$', stripped))


def _parse_table_cells(line: str) -> list[str]:
    """Splits a markdown table row line into clean text cell values."""
    parts = line.strip().strip('|').split('|')
    return [p.strip() for p in parts]


def parse_markdown_to_docx(doc: Document, markdown_text: str):
    """
    Parses full Markdown text line-by-line and renders it onto a python-docx Document.
    """
    lines = markdown_text.splitlines()
    idx = 0

    while idx < len(lines):
        line = lines[idx].strip()

        # Skip empty lines
        if not line:
            idx += 1
            continue

        # ---------------- 1. TABLE PARSER ---------------- #
        if _is_table_row(line):
            table_rows = []
            while idx < len(lines) and _is_table_row(lines[idx]):
                row_line = lines[idx].strip()
                if not _is_table_separator(row_line):
                    table_rows.append(_parse_table_cells(row_line))
                idx += 1

            if table_rows:
                headers = table_rows[0]
                data_rows = table_rows[1:]

                # Build Word Table
                table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
                
                # Header row text
                for c_idx, h_text in enumerate(headers):
                    table.rows[0].cells[c_idx].text = h_text

                # Data rows text
                for r_idx, row_data in enumerate(data_rows, start=1):
                    for c_idx, cell_val in enumerate(row_data):
                        if c_idx < len(headers):
                            table.rows[r_idx].cells[c_idx].text = cell_val

                # Format with AB Ark table styling
                format_data_table(table)
            continue

        # ---------------- 2. HEADING 1 (##) ---------------- #
        if line.startswith('## '):
            heading_text = line[3:].strip()
            add_custom_heading_1(doc, heading_text)

            # Check if this heading is a Phase Section (e.g. "3. Phase 1 — MVP Production Build")
            if "Phase 1" in heading_text or "PHASE 1" in heading_text:
                add_phase_banner(doc, "PHASE 1", "MVP Production Build", "Weeks 1 – 6", "$8,000")
            elif "Phase 2" in heading_text or "PHASE 2" in heading_text:
                add_phase_banner(doc, "PHASE 2", "V1 Production Platform", "Weeks 6 – 14", "$20,000")
            elif "Phase 3" in heading_text or "PHASE 3" in heading_text:
                add_phase_banner(doc, "PHASE 3", "Advanced Platform", "NA", "~")

            idx += 1
            continue

        # ---------------- 3. HEADING 2 (###) ---------------- #
        if line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=2)
            idx += 1
            continue

        # ---------------- 4. BULLET LIST (- or *) ---------------- #
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            clean_text = bullet_text.replace('**', '')
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(clean_text)
            idx += 1
            continue

        # ---------------- 5. REGULAR PARAGRAPH ---------------- #
        clean_para_text = line.replace('**', '')
        doc.add_paragraph(clean_para_text)
        idx += 1


# Verification runner when executed directly
if __name__ == "__main__":
    from docx import Document
    from .styles import register_proposal_styles

    test_md = """
## 1. Executive Summary

This proposal outlines the technical roadmap for Lumova AI.

### 1.1 Key Objectives

- High performance backend
- Multi-tenant data isolation

| Phase | Scope | Investment |
| --- | --- | --- |
| Phase 1 | Core MVP | $8,000 |
| Phase 2 | AI Features | $20,000 |
"""
    test_doc = Document()
    register_proposal_styles(test_doc)
    parse_markdown_to_docx(test_doc, test_md)
    print("--- markdown_parser.py verification completed successfully ---")
