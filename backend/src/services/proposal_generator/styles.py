
"""
styles.py — Word Style Definitions and Table Shading/Border Helpers for python-docx.

Registers custom paragraph and character styles onto a python-docx Document object,
ensuring consistent typography, heading bottom rules, and table formatting matching AB Ark's brand.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .constants import (
    FONT_PRIMARY, SIZE_H1, SIZE_H2, SIZE_BODY,
    RGB_PRIMARY_BLUE, RGB_TEXT_BLACK, RGB_WHITE,
    HEX_PRIMARY_BLUE, HEX_NAVY, HEX_BG_LIGHT_GRAY, HEX_BORDER_GRAY
)


# ---------------- XML HELPERS FOR TABLES & PARAGRAPHS ---------------- #

def set_cell_background(cell, hex_color: str):
    """Sets background shading hex color for a table cell using OpenXML."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (in twips) for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_bottom_border_to_paragraph(paragraph, hex_color: str = HEX_PRIMARY_BLUE, size_eighth_pt: int = 12):
    """Adds a thin horizontal bottom rule directly beneath a paragraph (e.g. for H1 headings)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighth_pt))  # 12 = 1.5pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def apply_table_borders(table, hex_color: str = HEX_BORDER_GRAY):
    """Applies subtle horizontal/vertical thin borders to an entire Word table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        tblBorders = OxmlElement('w:tblBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 0.5 pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), hex_color)
            tblBorders.append(border)
        tblPr[0].append(tblBorders)


# ---------------- STYLE REGISTRATION ---------------- #

def register_proposal_styles(doc: Document):
    """
    Registers standard AB Ark Word styles onto the Document.
    """
    styles = doc.styles

    # Default Normal style font
    normal_style = styles['Normal']
    normal_style.font.name = FONT_PRIMARY
    normal_style.font.size = SIZE_BODY
    normal_style.font.color.rgb = RGB_TEXT_BLACK

    # Heading 1 Style (Blue, 17pt, Bold)
    try:
        h1_style = styles['Heading 1']
    except KeyError:
        h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = FONT_PRIMARY
    h1_style.font.size = SIZE_H1
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGB_PRIMARY_BLUE
    h1_style.paragraph_format.space_before = Pt(16)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.keep_with_next = True

    # Heading 2 Style (Black, 13pt, Bold)
    try:
        h2_style = styles['Heading 2']
    except KeyError:
        h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = FONT_PRIMARY
    h2_style.font.size = SIZE_H2
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGB_TEXT_BLACK
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(4)
    h2_style.paragraph_format.keep_with_next = True

    # List Bullet Style
    try:
        bullet_style = styles['List Bullet']
    except KeyError:
        bullet_style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = FONT_PRIMARY
    bullet_style.font.size = SIZE_BODY
    bullet_style.font.color.rgb = RGB_TEXT_BLACK
    bullet_style.paragraph_format.space_before = Pt(2)
    bullet_style.paragraph_format.space_after = Pt(3)
    bullet_style.paragraph_format.left_indent = Inches(0.25)


def add_custom_heading_1(doc: Document, text: str):
    """
    Adds Heading 1 with automatic thin blue rule beneath it.
    """
    h = doc.add_paragraph(text, style='Heading 1')
    add_bottom_border_to_paragraph(h, hex_color=HEX_PRIMARY_BLUE, size_eighth_pt=12)
    return h


def format_data_table(table, col_widths: list = None):
    """
    Formats a standard python-docx table with AB Ark styling:
    - Dark Navy header row with bold white text
    - Alternating row zebra shading (#F8FAFC)
    - Thin borders and padded cells
    """
    apply_table_borders(table, hex_color=HEX_BORDER_GRAY)

    # Format Header Row
    header_row = table.rows[0]
    for cell in header_row.cells:
        set_cell_background(cell, HEX_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = FONT_PRIMARY
                r.font.bold = True
                r.font.size = SIZE_BODY
                r.font.color.rgb = RGB_WHITE

    # Format Data Rows with Alternating Shading
    for row_idx, row in enumerate(table.rows[1:], start=1):
        bg_color = HEX_BG_LIGHT_GRAY if row_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = FONT_PRIMARY
                    r.font.size = SIZE_BODY
                    r.font.color.rgb = RGB_TEXT_BLACK

    # Set optional column widths
    if col_widths:
        for row in table.rows:
            for cell, width in zip(row.cells, col_widths):
                cell.width = width


# Verification runner when executed directly
if __name__ == "__main__":
    doc = Document()
    register_proposal_styles(doc)
    add_custom_heading_1(doc, "1. Executive Summary Test")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Header 1"
    t.rows[0].cells[1].text = "Header 2"
    t.rows[1].cells[0].text = "Data 1"
    t.rows[1].cells[1].text = "Data 2"
    format_data_table(t)
    print("--- styles.py verification completed successfully ---")
