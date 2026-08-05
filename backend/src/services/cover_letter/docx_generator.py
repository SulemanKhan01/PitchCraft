"""
docx_generator.py — Minimalist, text-only DOCX generator for cover letters.

Mirrors generate_minimal_pdf:
- Margins: 1 inch
- Font: 11pt Calibri
- Paragraph spacing preserved (one paragraph per line)
- Bullet lines ("- ", "* ", "•") -> List Bullet style
- Bold segments (**text**) rendered bold
"""

import io
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING


def _split_bold_segments(text: str) -> list[tuple[str, bool]]:
    """Splits text into (segment, is_bold) pairs for **bold** markers."""
    segments = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], False))
        segments.append((match.group(1), True))
        last_end = match.end()
    if last_end < len(text):
        segments.append((text[last_end:], False))
    return segments or [(text, False)]


def generate_minimal_docx(cover_letter_text: str) -> bytes:
    """
    Generates a minimalist, text-only DOCX in-memory.
    - Margins: 1 inch
    - Font: 11pt Calibri
    - Color: Charcoal gray (#2D3748)
    - Space after paragraphs: 14pt
    """
    buffer = io.BytesIO()

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    pf = normal.paragraph_format
    pf.space_after = Pt(14)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.45

    raw_paragraphs = [p.strip() for p in cover_letter_text.split('\n') if p.strip()]

    for p_text in raw_paragraphs:
        is_bullet = p_text.startswith('- ') or p_text.startswith('* ') or p_text.startswith('•')
        if is_bullet:
            clean_text = p_text[2:] if p_text.startswith('- ') or p_text.startswith('* ') else p_text[1:]
            p = doc.add_paragraph(style='List Bullet')
        else:
            clean_text = p_text
            p = doc.add_paragraph()

        for segment, is_bold in _split_bold_segments(clean_text):
            run = p.add_run(segment)
            if is_bold:
                run.bold = True

    doc.save(buffer)

    docx_bytes = buffer.getvalue()
    buffer.close()

    return docx_bytes
